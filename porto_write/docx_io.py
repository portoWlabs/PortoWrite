import logging
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from porto_write.document import PortoDocument, Chapter, TextBlock

logger = logging.getLogger(__name__)

# Mapping PortoWrite style names to Word built-in style names
PW_TO_WORD_STYLES = {
    "ChapterHeader": "Heading 1",
    "SubHeader": "Heading 2",
    "Body": "Normal",
    "BlockQuote": "Quote"
}

# Reverse mapping for import
WORD_TO_PW_STYLES = {v: k for k, v in PW_TO_WORD_STYLES.items()}

def export_docx(doc: PortoDocument, file_path: str):
    """Export a PortoDocument to a Word (.docx) file."""
    word_doc = Document()
    
    # 1. Metadata
    word_doc.core_properties.title = doc.title
    word_doc.core_properties.author = doc.author
    
    # 2. Chapters and Blocks
    for i, chapter in enumerate(doc.chapters):
        # We handle chapter titles as Heading 1
        if chapter.title:
            word_doc.add_paragraph(chapter.title, style='Heading 1')
            
        for block in chapter.blocks:
            word_style = PW_TO_WORD_STYLES.get(block.style_name, "Normal")
            
            # Ensure the style exists in the default template, fallback to Normal
            try:
                p = word_doc.add_paragraph(block.text, style=word_style)
            except KeyError:
                p = word_doc.add_paragraph(block.text, style="Normal")
            
            # Note: Inline formatting (bold, italic, underline) inside TextBlock 
            # is not yet supported by our flat TextBlock model. 
            # Once we support spans, we would apply them here.
            
    word_doc.save(file_path)
    logger.info("Exported Word (.docx) to %s", file_path)

def import_docx(file_path: str) -> PortoDocument:
    """Import a Word (.docx) file into a PortoDocument."""
    word_doc = Document(file_path)
    
    doc = PortoDocument(
        title=word_doc.core_properties.title or "Untitled",
        author=word_doc.core_properties.author or ""
    )
    
    current_chapter = None
    
    for para in word_doc.paragraphs:
        if not para.text.strip():
            continue
            
        style_name = para.style.name
        pw_style = WORD_TO_PW_STYLES.get(style_name, "Body")
        
        # If it's a Heading 1, start a new chapter
        if style_name == "Heading 1":
            current_chapter = doc.add_chapter(para.text)
        else:
            if not current_chapter:
                current_chapter = doc.add_chapter("Chapter")
            current_chapter.add_block(pw_style, para.text)
            
    logger.info("Imported Word (.docx) from %s", file_path)
    return doc
